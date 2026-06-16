#!/usr/bin/env python3
"""
Convert progress report markdown to well-structured Word document.
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed. Please run: pip3 install python-docx")
    exit(1)

import re
from pathlib import Path


def parse_markdown(md_text):
    """Parse markdown text into structured sections."""
    lines = md_text.split('\n')
    sections = []
    current_section = None
    current_content = []
    
    for line in lines:
        # Main title (# Title)
        if line.startswith('# ') and not line.startswith('## '):
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {
                'type': 'title',
                'level': 1,
                'text': line[2:].strip(),
                'content': ''
            }
            current_content = []
        # Section headers (## Header)
        elif line.startswith('## ') and not line.startswith('### '):
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {
                'type': 'heading',
                'level': 2,
                'text': line[3:].strip(),
                'content': ''
            }
            current_content = []
        # Subsection headers (### Header)
        elif line.startswith('### '):
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {
                'type': 'heading',
                'level': 3,
                'text': line[4:].strip(),
                'content': ''
            }
            current_content = []
        # Horizontal rule
        elif line.strip() == '---':
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
                current_content = []
            sections.append({'type': 'separator'})
            current_section = None
        # Content
        else:
            current_content.append(line)
    
    # Add last section
    if current_section:
        current_section['content'] = '\n'.join(current_content)
        sections.append(current_section)
    
    return sections


def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False, size=None):
    """Add a paragraph with formatting."""
    # Handle bold markdown
    text = text.strip()
    if not text:
        doc.add_paragraph()
        return
    
    paragraph = doc.add_paragraph(style=style)
    
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Normal text
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        
        if size:
            run.font.size = Pt(size)


def process_content(doc, content):
    """Process content and add to document with formatting."""
    if not content.strip():
        return
    
    lines = content.split('\n')
    in_list = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            if in_list:
                in_list = False
            continue
        
        # Bullet point
        if line.startswith('- '):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')
            in_list = True
        # Italic text (at start of line, e.g., conclusions)
        elif line.startswith('*') and line.endswith('*') and line.count('*') == 2:
            add_formatted_paragraph(doc, line[1:-1], italic=True)
            in_list = False
        # Bold subheading in content
        elif line.endswith(':') and '**' in line:
            add_formatted_paragraph(doc, line, bold=True, size=11)
            in_list = False
        # Normal paragraph
        else:
            add_formatted_paragraph(doc, line)
            in_list = False


def create_word_document(md_file_path, output_path):
    """Convert markdown file to Word document."""
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    # Parse markdown
    sections = parse_markdown(md_text)
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading styles
    heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(18)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 102, 153)
    
    heading2_style = styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Calibri'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 102, 153)
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # Process sections
    for section in sections:
        if section['type'] == 'title':
            # Main title
            title = doc.add_heading(section['text'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Add space after title
            
        elif section['type'] == 'heading':
            # Add spacing before major sections
            if section['level'] == 2:
                doc.add_paragraph()
            
            # Section heading
            doc.add_heading(section['text'], level=section['level'])
            
            # Process content
            if section.get('content'):
                process_content(doc, section['content'])
        
        elif section['type'] == 'separator':
            # Add horizontal line
            doc.add_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('_' * 80)
            run.font.color.rgb = RGBColor(192, 192, 192)
            doc.add_paragraph()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Save document
    doc.save(output_path)
    print(f"✓ Word document created: {output_path}")


if __name__ == '__main__':
    script_dir = Path(__file__).parent
    md_file = script_dir / 'progress_report.md'
    output_file = script_dir / 'progress_report.docx'
    
    if not md_file.exists():
        print(f"ERROR: Markdown file not found: {md_file}")
        exit(1)
    
    create_word_document(md_file, output_file)
    print(f"✓ Conversion complete!")
    print(f"  Input:  {md_file}")
    print(f"  Output: {output_file}")
